"""Создание Word-документа с фотографиями."""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Cm, Mm
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from .config import DocumentConfig
from .image_loader import get_image_dimensions
from .utils import (
    get_grid_dimensions,
    calculate_image_size,
    set_cell_margins,
    add_page_number_field,
)

logger = logging.getLogger(__name__)

# A4 portrait (210×297) — базовая ориентация Word, затем swap для landscape
A4_PORTRAIT_WIDTH_MM = 210
A4_PORTRAIT_HEIGHT_MM = 297
DEFAULT_MARGIN_MM = 10  # поля страницы


class PhotoDocumentBuilder:
    """Создаёт Word-документ с фотографиями в сетке."""

    def __init__(self, config: DocumentConfig):
        self.config = config
        self.document = Document()
        self._setup_first_section()
        self._page_count = 0

    def _setup_first_section(self) -> None:
        """Настраивает первую секцию: альбомная ориентация, разный футер на первой странице."""
        section = self.document.sections[0]
        section.page_width = Mm(A4_PORTRAIT_WIDTH_MM)
        section.page_height = Mm(A4_PORTRAIT_HEIGHT_MM)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        section.left_margin = Mm(DEFAULT_MARGIN_MM)
        section.right_margin = Mm(DEFAULT_MARGIN_MM)
        section.top_margin = Mm(DEFAULT_MARGIN_MM)
        section.bottom_margin = Mm(DEFAULT_MARGIN_MM)

        section.different_first_page_header_footer = True

    def _get_content_dimensions(self) -> tuple[float, float]:
        """Возвращает (width_cm, height_cm) доступной области для контента."""
        section = self.document.sections[0]
        page_width_cm = section.page_width.cm
        page_height_cm = section.page_height.cm
        left = section.left_margin.cm
        right = section.right_margin.cm
        top = section.top_margin.cm
        bottom = section.bottom_margin.cm
        return page_width_cm - left - right, page_height_cm - top - bottom

    def _add_page_number_footer(self, section) -> None:
        """Добавляет нумерацию страниц в футер (для страниц 2+)."""
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = 1  # CENTER
        run = paragraph.add_run()
        add_page_number_field(run, "PAGE")

    def add_page_with_photos(self, photo_paths: list[Path]) -> None:
        """Добавляет страницу с фотографиями в сетке."""
        if not photo_paths:
            return

        rows, cols = get_grid_dimensions(self.config.photos_per_page, photo_paths)
        content_width_cm, content_height_cm = self._get_content_dimensions()

        cell_width_cm = content_width_cm / cols
        cell_height_cm = content_height_cm / rows

        margin_per_side_cm = self.config.margin_cm / 2
        available_width_cm = cell_width_cm - self.config.margin_cm
        available_height_cm = cell_height_cm - self.config.margin_cm

        if self._page_count > 0:
            self.document.add_section(WD_SECTION.NEW_PAGE)
            new_section = self.document.sections[-1]
            new_section.page_width = Mm(A4_PORTRAIT_WIDTH_MM)
            new_section.page_height = Mm(A4_PORTRAIT_HEIGHT_MM)
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width, new_section.page_height = (
                new_section.page_height,
                new_section.page_width,
            )
            new_section.left_margin = Mm(DEFAULT_MARGIN_MM)
            new_section.right_margin = Mm(DEFAULT_MARGIN_MM)
            new_section.top_margin = Mm(DEFAULT_MARGIN_MM)
            new_section.bottom_margin = Mm(DEFAULT_MARGIN_MM)
            new_section.different_first_page_header_footer = True
            self._add_page_number_footer(new_section)

        table = self.document.add_table(rows=rows, cols=cols)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(cell_height_cm)

        photo_index = 0
        for row in table.rows:
            for cell in row.cells:
                if photo_index < len(photo_paths):
                    path = photo_paths[photo_index]
                    set_cell_margins(
                        cell,
                        top=margin_per_side_cm,
                        start=margin_per_side_cm,
                        bottom=margin_per_side_cm,
                        end=margin_per_side_cm,
                    )
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_before = 0
                    paragraph.paragraph_format.space_after = 0
                    paragraph.alignment = 1  # CENTER

                    dims = get_image_dimensions(path)
                    if dims:
                        img_w, img_h = dims
                        width_cm, height_cm = calculate_image_size(
                            available_width_cm,
                            available_height_cm,
                            img_w,
                            img_h,
                        )
                        run = paragraph.add_run()
                        try:
                            run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
                        except Exception as e:
                            logger.error("Не удалось вставить %s: %s", path.name, e)
                    photo_index += 1

        self._page_count += 1

    def save(self, path: Path | str) -> None:
        """Сохраняет документ."""
        self.document.save(str(path))
